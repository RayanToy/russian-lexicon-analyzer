# src/highlighting/highlighter.py
from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path
from typing import List, Optional, Set

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.highlighting.lexicon import (
    LemmaNormalizer,
    LexiconRepository,
    load_stopword_file,
)
from src.highlighting.matcher import TokenLemmaMatcher
from src.preprocessing.text_preprocessor import StandardTextPreprocessor


class DocxHighlighter:
    """
    Подсвечивает слова в DOCX-документе по результатам классификации.

    Цвета по умолчанию:
        unknown    → красный (FF0000)
        stop_black → чёрный (000000)
        known      → без изменений (наследует из стиля)
        stopword   → без изменений

    Корректно обрабатывает:
        - soft hyphen (U+00AD) внутри слов
        - слова, разбитые на несколько run'ов
        - таблицы и вложенные таблицы
        - форматирование (жирный, курсив, размер шрифта)
    """

    # Паттерн для токенизации: буквы Unicode + ударение,
    # дефисные конструкции (дефис внутри слова)
    _LETTER  = r"[^\W\d_]"
    _STRESS  = r"[\u0301\u0341\u00B4\u02CA]"
    WORD_RE  = re.compile(
        rf"(?:{_LETTER}|{_STRESS})+"
        rf"(?:[\u00AD‐‑–—-](?:{_LETTER}|{_STRESS})+)*",
        re.UNICODE,
    )

    def __init__(
        self,
        preprocessor: StandardTextPreprocessor,
        lexicon_lemmas: Set[str],
        *,
        unknown_hex: str = "FF0000",
        black_hex:   str = "000000",
        stopword_lemmas:       Optional[Set[str]] = None,
        black_stopword_lemmas: Optional[Set[str]] = None,
        strip_prefixes:        Optional[Set[str]] = None,
    ) -> None:
        """
        Args:
            preprocessor: препроцессор текста
            lexicon_lemmas: множество «известных» лемм
            unknown_hex: цвет для незнакомых слов (RGB hex, без #)
            black_hex: цвет для слов из чёрного списка
            stopword_lemmas: стоп-слова (не подсвечивать)
            black_stopword_lemmas: слова из чёрного списка
            strip_prefixes: приставки для снятия при поиске в лексиконе
        """
        self.unknown_hex = unknown_hex
        self.black_hex   = black_hex
        self.matcher     = TokenLemmaMatcher(
            preprocessor,
            lexicon_lemmas,
            stopword_lemmas=stopword_lemmas,
            black_stopword_lemmas=black_stopword_lemmas,
            strip_prefixes=strip_prefixes,
        )

    # ------------------------------------------------------------------
    # Публичный метод
    # ------------------------------------------------------------------
    def highlight_file(self, in_path: str | Path, out_path: str | Path) -> None:
        """
        Читает DOCX, подсвечивает слова, сохраняет результат.

        Args:
            in_path: путь к исходному DOCX
            out_path: путь для сохранения результата
        """
        doc = Document(str(in_path))
        for paragraph in self._iter_all_paragraphs(doc):
            self._process_paragraph(paragraph)
        doc.save(str(out_path))

    # ------------------------------------------------------------------
    # Обход документа
    # ------------------------------------------------------------------
    def _iter_all_paragraphs(self, doc: Document):
        """Итерирует по всем абзацам документа, включая таблицы."""
        yield from doc.paragraphs
        for table in doc.tables:
            yield from self._iter_paragraphs_in_table(table)

    def _iter_paragraphs_in_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested_table in cell.tables:
                    yield from self._iter_paragraphs_in_table(nested_table)

    # ------------------------------------------------------------------
    # Обработка абзаца
    # ------------------------------------------------------------------
    def _process_paragraph(self, paragraph) -> None:
        """
        Обрабатывает абзац:
        - передаёт полный текст абзаца в NER (для контекста)
        - обрабатывает каждый run
        - склеивает слова на границах run'ов
        """
        try:
            self.matcher._ner_context = paragraph.text
        except Exception:
            self.matcher._ner_context = None

        runs = list(paragraph.runs)
        if not runs:
            return

        # Обрабатываем первый run
        prev_run    = runs[0]
        prev_pieces = self._collect_pieces(prev_run._element)
        self._apply_intra_run_joins(prev_pieces)

        for cur_run in runs[1:]:
            cur_pieces = self._collect_pieces(cur_run._element)
            self._apply_intra_run_joins(cur_pieces)

            # Склейка на границе двух run'ов
            self._apply_cross_run_join(prev_pieces, cur_pieces)

            # Перестраиваем предыдущий run если есть совпадения
            if self._any_matched(prev_pieces):
                self._rebuild_run(prev_run, prev_pieces)

            prev_run, prev_pieces = cur_run, cur_pieces

        # Последний run
        if self._any_matched(prev_pieces):
            self._rebuild_run(prev_run, prev_pieces)

    # ------------------------------------------------------------------
    # Сбор «кусков» run'а
    # ------------------------------------------------------------------
    def _collect_pieces(self, r_el) -> List[dict]:
        """
        Разбирает XML-элемент run на «куски»:
        - 'text'    → w:t (текст)
        - 'control' → w:softHyphen, w:noBreakHyphen, w:br, w:tab
        - 'other'   → всё остальное
        """
        pieces = []
        for child in list(r_el):
            name = child.tag.split('}', 1)[-1] if '}' in child.tag else child.tag
            if name == "t":
                text     = child.text or ""
                segments = self._tokenize_text(text)
                pieces.append({"type": "text", "segments": segments, "orig_el": child})
            elif name in ("softHyphen", "noBreakHyphen", "br", "tab"):
                pieces.append({"type": "control", "name": name, "orig_el": child})
            else:
                pieces.append({"type": "other", "orig_el": child})
        return pieces

    def _tokenize_text(self, text: str) -> List[dict]:
        """
        Токенизирует строку на слова и «не-слова».
        Каждый сегмент — dict с полями:
            text, is_word, matched, category
        """
        segments: List[dict] = []
        last = 0
        for m in self.WORD_RE.finditer(text):
            start, end = m.start(), m.end()
            if start > last:
                segments.append({
                    "text": text[last:start],
                    "is_word": False, "matched": False, "category": None,
                })
            token    = text[start:end]
            category = self.matcher.classify(token)
            matched  = category in ('unknown', 'stop_black')
            segments.append({
                "text": token,
                "is_word": True, "matched": matched, "category": category,
            })
            last = end
        if last < len(text):
            segments.append({
                "text": text[last:],
                "is_word": False, "matched": False, "category": None,
            })
        return segments

    # ------------------------------------------------------------------
    # Склейки
    # ------------------------------------------------------------------
    def _last_word_seg_idx(self, piece: dict) -> Optional[int]:
        """
        Индекс последнего словесного сегмента в piece,
        если после него нет никаких символов (даже пробела).
        """
        if piece["type"] != "text":
            return None
        segs = piece["segments"]
        for j in range(len(segs) - 1, -1, -1):
            if segs[j]["is_word"]:
                trailing = ''.join(s["text"] for s in segs[j + 1:])
                return j if trailing == "" else None
        return None

    def _first_word_seg_idx(self, piece: dict) -> Optional[int]:
        """
        Индекс первого словесного сегмента в piece,
        если перед ним нет никаких символов.
        """
        if piece["type"] != "text":
            return None
        segs = piece["segments"]
        for j, seg in enumerate(segs):
            if seg["is_word"]:
                prefix = ''.join(s["text"] for s in segs[:j])
                return j if prefix == "" else None
        return None

    def _mark_pair(
        self,
        left_seg: dict,
        right_seg: dict,
        category: str,
    ) -> None:
        """Помечает пару сегментов одной категорией."""
        matched = category in ('unknown', 'stop_black')
        for seg in (left_seg, right_seg):
            seg["matched"]  = matched
            seg["category"] = category

    def _apply_intra_run_joins(self, pieces: List[dict]) -> None:
        """
        Склеивает слова внутри run'а через softHyphen/noBreakHyphen.
        Также обрабатывает смежные text-куски без разделителя.
        """
        i = 0
        while i < len(pieces) - 2:
            a, b, c = pieces[i], pieces[i + 1], pieces[i + 2]
            if (
                a["type"] == "text"
                and b["type"] == "control"
                and b["name"] in ("softHyphen", "noBreakHyphen")
                and c["type"] == "text"
            ):
                li = self._last_word_seg_idx(a)
                ri = self._first_word_seg_idx(c)
                if li is not None and ri is not None:
                    left_tok  = a["segments"][li]["text"]
                    right_tok = c["segments"][ri]["text"]
                    category  = self.matcher.classify_hyphenated([left_tok, right_tok])
                    self._mark_pair(a["segments"][li], c["segments"][ri], category)
                i += 3
            elif a["type"] == "text" and b["type"] == "text":
                # Смежные text-куски (без разделителя)
                li = self._last_word_seg_idx(a)
                ri = self._first_word_seg_idx(b)
                if li is not None and ri is not None:
                    combined = a["segments"][li]["text"] + b["segments"][ri]["text"]
                    category = self.matcher.classify(combined)
                    self._mark_pair(a["segments"][li], b["segments"][ri], category)
                i += 1
            else:
                i += 1

    def _apply_cross_run_join(
        self,
        prev_pieces: List[dict],
        cur_pieces:  List[dict],
    ) -> None:
        """
        Склеивает слово на границе двух run'ов.
        Работает только если между словами нет пробелов.
        """
        # Находим последний словесный сегмент предыдущего run'а
        left_piece = left_idx = None
        for p in reversed(prev_pieces):
            if p["type"] == "text":
                li = self._last_word_seg_idx(p)
                if li is not None:
                    left_piece, left_idx = p, li
                    break

        # Находим первый словесный сегмент текущего run'а
        right_piece = right_idx = None
        for p in cur_pieces:
            if p["type"] == "text":
                ri = self._first_word_seg_idx(p)
                if ri is not None:
                    right_piece, right_idx = p, ri
                    break

        if left_piece is None or right_piece is None:
            return

        left_token  = left_piece["segments"][left_idx]["text"]
        right_token = right_piece["segments"][right_idx]["text"]
        combined    = left_token + right_token
        category    = self.matcher.classify(combined)
        self._mark_pair(
            left_piece["segments"][left_idx],
            right_piece["segments"][right_idx],
            category,
        )

    # ------------------------------------------------------------------
    # Перестройка run'а
    # ------------------------------------------------------------------
    def _any_matched(self, pieces: List[dict]) -> bool:
        return any(
            piece["type"] == "text"
            and any(seg.get("matched") for seg in piece["segments"])
            for piece in pieces
        )

    def _rebuild_run(self, run, pieces: List[dict]) -> None:
        """
        Перестраивает run: разбивает на отдельные run'ы по сегментам,
        назначая нужный цвет каждому.
        """
        r_el   = run._element
        parent = r_el.getparent()
        prev_r = r_el

        for piece in pieces:
            ptype = piece.get("type")

            if ptype == "text":
                for seg in piece["segments"]:
                    txt = seg.get("text", "")
                    if not txt:
                        continue

                    new_r = self._clone_rpr(r_el)
                    self._append_text(new_r, txt)

                    if seg.get("is_word") and seg.get("matched"):
                        color = (
                            self.black_hex
                            if seg.get("category") == 'stop_black'
                            else self.unknown_hex
                        )
                        self._set_color(new_r, color)
                    else:
                        self._clear_color(new_r)

                    prev_r.addnext(new_r)
                    prev_r = new_r

            else:
                # control / other — копируем как есть, цвет не трогаем
                new_r = self._clone_rpr(r_el)
                new_r.append(deepcopy(piece["orig_el"]))
                self._clear_color(new_r)
                prev_r.addnext(new_r)
                prev_r = new_r

        parent.remove(r_el)

    # ------------------------------------------------------------------
    # XML-хелперы
    # ------------------------------------------------------------------
    def _clone_rpr(self, src_r) -> OxmlElement:
        """Создаёт новый run с копией rPr (форматирование) исходного."""
        new_r = OxmlElement("w:r")
        if src_r.rPr is not None:
            new_r.append(deepcopy(src_r.rPr))
        return new_r

    def _append_text(self, r_el, text: str) -> None:
        """Добавляет w:t с нужным атрибутом xml:space."""
        t = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            t.set(qn("xml:space"), "preserve")
        t.text = text
        r_el.append(t)

    def _set_color(self, r_el, hex_color: str) -> None:
        """Устанавливает явный цвет шрифта в rPr."""
        rPr = self._ensure_rpr(r_el)
        color = rPr.find("./w:color", namespaces=r_el.nsmap)
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), hex_color)

    def _clear_color(self, r_el) -> None:
        """Убирает явный цвет — текст наследует цвет из стиля."""
        rPr = r_el.find("./w:rPr", namespaces=r_el.nsmap)
        if rPr is None:
            return
        color = rPr.find("./w:color", namespaces=r_el.nsmap)
        if color is not None:
            rPr.remove(color)

    def _ensure_rpr(self, r_el) -> OxmlElement:
        """Возвращает rPr элемент, создавая его если нет."""
        rPr = r_el.find("./w:rPr", namespaces=r_el.nsmap)
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            r_el.insert(0, rPr)
        return rPr


# ---------------------------------------------------------------------------
# Публичная функция-обёртка
# ---------------------------------------------------------------------------
def highlight_docx(
    in_docx: str | Path,
    out_docx: str | Path,
    lexicon_lemmas: Set[str],
    *,
    unknown_color_hex: str = "FF0000",
    black_stopwords_path: Optional[str | Path] = None,
    strip_prefixes: Optional[Set[str]] = None,
) -> None:
    """
    Подсвечивает незнакомые слова в DOCX-документе.

    Args:
        in_docx: исходный DOCX-файл
        out_docx: путь для сохранения результата
        lexicon_lemmas: множество «известных» лемм
        unknown_color_hex: цвет для незнакомых слов (RGB hex без #)
        black_stopwords_path: путь к файлу с чёрным списком (одно слово на строку)
        strip_prefixes: приставки для снятия при поиске в лексиконе

    Example:
        >>> from src.highlighting.lexicon import LexiconRepository
        >>> lemmas = LexiconRepository().load_lemmas("data/lexicon.xlsx")
        >>> highlight_docx("input.docx", "output.docx", lemmas)
    """
    pre            = StandardTextPreprocessor()
    black_stopwords = (
        load_stopword_file(black_stopwords_path)
        if black_stopwords_path else set()
    )
    highlighter = DocxHighlighter(
        pre,
        lexicon_lemmas,
        unknown_hex=unknown_color_hex,
        black_hex="000000",
        black_stopword_lemmas=black_stopwords,
        strip_prefixes=strip_prefixes,
    )
    highlighter.highlight_file(in_docx, out_docx)
    print(f"Готово: {out_docx}")