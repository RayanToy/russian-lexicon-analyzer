# src/analysis/file_processor.py
from __future__ import annotations

import os
import re
from pathlib import Path
from typing import List, Optional

import docx
import pandas as pd
from docx import Document
from docx.shared import RGBColor

from src.preprocessing.text_preprocessor import TextPreprocessor
from src.analysis.frequency_analyzer import FrequencyAnalyzer, FreqEntry


class FileProcessor:
    """
    Обрабатывает DOCX-файлы: извлекает текст, лемматизирует,
    строит частотные списки и сохраняет результаты в XLSX.

    Текст обрабатывается блоками (block_size_threshold слов),
    чтобы не загружать в память весь документ целиком.

    Соглашение по именованию файлов:
        <любые_части>_<кол-во_слов>.docx
        Например: corpus_history_10_1750000.docx
        Последний числовой сегмент перед .docx — объём корпуса в словах.
    """

    def __init__(
        self,
        preprocessor: TextPreprocessor,
        frequency_analyzer: FrequencyAnalyzer,
        block_size_threshold: int = 1_000_000,
    ) -> None:
        """
        Args:
            preprocessor: экземпляр препроцессора текста
            frequency_analyzer: экземпляр анализатора частотности
            block_size_threshold: макс. кол-во слов в одном блоке обработки
        """
        self.preprocessor         = preprocessor
        self.frequency_analyzer   = frequency_analyzer
        self.block_size_threshold = block_size_threshold

    # ------------------------------------------------------------------
    # Внутренние методы
    # ------------------------------------------------------------------
    def _extract_total_words(self, filename: str) -> int:
        """
        Извлекает объём корпуса из имени файла.
        Ожидается формат: <...>_<число>.docx

        Returns:
            Кол-во слов или 0, если число не найдено.
        """
        match = re.search(r'_(\d+)\.docx$', filename, flags=re.IGNORECASE)
        return int(match.group(1)) if match else 0

    def _iter_paragraph_texts(self, docx_path: Path) -> List[str]:
        """Извлекает тексты абзацев из DOCX-файла."""
        doc = docx.Document(str(docx_path))
        return [para.text for para in doc.paragraphs if para.text.strip()]

    def _process_block(
        self,
        texts: List[str],
        total_words: int,
        skip_proper_nouns: bool,
    ) -> List[FreqEntry]:
        """Лемматизирует блок текстов и возвращает частотный список."""
        combined = ' '.join(texts)
        lemmas   = self.preprocessor.preprocess(
            combined,
            skip_proper_nouns=skip_proper_nouns,
        )
        return self.frequency_analyzer.analyze_frequency(lemmas, total_words)

    # ------------------------------------------------------------------
    # Основные методы
    # ------------------------------------------------------------------
    def process_text_in_blocks(
        self,
        docx_path: Path | str,
        skip_proper_nouns: bool = True,
    ) -> List[List[FreqEntry]]:
        """
        Читает DOCX и обрабатывает текст блоками по block_size_threshold слов.

        Args:
            docx_path: путь к DOCX-файлу
            skip_proper_nouns: пропускать ли собственные имена при лемматизации

        Returns:
            Список блоков, каждый блок — частотный список [(слово, abs, ipm), ...]
        """
        docx_path   = Path(docx_path)
        total_words = self._extract_total_words(docx_path.name)

        paragraphs    = self._iter_paragraph_texts(docx_path)
        freq_dist_list: List[List[FreqEntry]] = []
        block_texts:   List[str] = []
        block_word_count = 0

        for para_text in paragraphs:
            # Нормализуем текст абзаца
            text  = ' ' + para_text.lower().replace('.', ' ')
            words = text.split()

            if block_word_count + len(words) <= self.block_size_threshold:
                block_texts.append(text)
                block_word_count += len(words)
            else:
                # Блок заполнен — обрабатываем и начинаем новый
                if block_texts:
                    fdist = self._process_block(block_texts, total_words, skip_proper_nouns)
                    freq_dist_list.append(fdist)
                block_texts      = [text]
                block_word_count = len(words)

        # Остаток
        if block_texts:
            fdist = self._process_block(block_texts, total_words, skip_proper_nouns)
            freq_dist_list.append(fdist)

        return freq_dist_list

    def create_frequency_lists(
        self,
        folder_path: Path | str,
        skip_proper_nouns: bool = True,
        output_folder: Optional[Path | str] = None,
    ) -> None:
        """
        Обрабатывает все DOCX-файлы в папке и сохраняет частотные списки в XLSX.

        Args:
            folder_path: папка с DOCX-файлами
            skip_proper_nouns: пропускать ли собственные имена
            output_folder: куда сохранять XLSX.
                           Если None — сохраняется рядом с DOCX.
        """
        folder_path   = Path(folder_path)
        output_folder = Path(output_folder) if output_folder else folder_path
        output_folder.mkdir(parents=True, exist_ok=True)

        docx_files = list(folder_path.glob('*.docx'))
        if not docx_files:
            print(f"DOCX-файлы не найдены в: {folder_path}")
            return

        for docx_path in docx_files:
            print(f"  Обработка: {docx_path.name}")
            freq_dist_list = self.process_text_in_blocks(
                docx_path,
                skip_proper_nouns=skip_proper_nouns,
            )

            # Сворачиваем все блоки в один DataFrame
            data = [
                [word, abs_freq, norm_freq]
                for block in freq_dist_list
                for word, abs_freq, norm_freq in block
            ]
            df = pd.DataFrame(
                data,
                columns=['Слово', 'Сумма слов', 'Нормализованная частота'],
            )

            output_path = output_folder / (docx_path.stem + '.xlsx')
            df.to_excel(output_path, index=False)
            print(f"  Сохранено: {output_path}")

        print(f"\nГотово. Результаты в: {output_folder}")

    def create_frequency_lists_recursively(
        self,
        root_folder: Path | str,
        output_folder: Optional[Path | str] = None,
        skip_proper_nouns: bool = False,
    ) -> None:
        """
        Рекурсивно обходит подпапки и обрабатывает все DOCX-файлы.
        Структура подпапок воспроизводится в output_folder.

        Args:
            root_folder: корневая папка
            output_folder: куда сохранять результаты (структура сохраняется).
                           Если None — файлы сохраняются рядом с исходными.
            skip_proper_nouns: пропускать ли собственные имена
        """
        root_folder = Path(root_folder)

        for dirpath, _, filenames in os.walk(root_folder):
            dirpath = Path(dirpath)
            if not any(f.endswith('.docx') for f in filenames):
                continue

            print(f"Обработка папки: {dirpath}")

            # Вычисляем выходную папку с сохранением структуры
            if output_folder is not None:
                relative         = dirpath.relative_to(root_folder)
                current_out      = Path(output_folder) / relative
            else:
                current_out = None

            self.create_frequency_lists(
                dirpath,
                skip_proper_nouns=skip_proper_nouns,
                output_folder=current_out,
            )

        print(f"\nОбработка завершена. Корневая папка: {root_folder}")

    def clean_docx(
        self,
        input_path: Path | str,
        output_path: Path | str,
    ) -> None:
        """
        Очищает DOCX-файлы: удаляет изображения, конвертирует таблицы в текст,
        убирает спецсимволы. Результат сохраняется с префиксом 'cleaned_'.

        Args:
            input_path: папка с исходными DOCX-файлами
            output_path: папка для сохранения очищенных файлов
        """
        input_path  = Path(input_path)
        output_path = Path(output_path)

        if not input_path.is_dir():
            raise ValueError(f"Путь должен быть папкой: {input_path}")

        output_path.mkdir(parents=True, exist_ok=True)

        for docx_file in input_path.glob('*.docx'):
            output_file = output_path / f"cleaned_{docx_file.name}"
            try:
                self._clean_single_docx(docx_file, output_file)
                print(f"Обработан: {docx_file.name} → {output_file.name}")
            except Exception as e:
                print(f"Ошибка при обработке {docx_file.name}: {e}")

    def _clean_single_docx(self, input_file: Path, output_file: Path) -> None:
        """Очищает один DOCX-файл."""
        doc     = Document(str(input_file))
        new_doc = Document()

        # Сначала конвертируем таблицы в текст
        table_texts: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cleaned = self.preprocessor.remove_special_chars_docx(
                        cell.text.strip()
                    )
                    if cleaned:
                        row_text.append(cleaned)
                if row_text:
                    sentence = " ".join(row_text)
                    if not sentence.endswith('.'):
                        sentence += '.'
                    table_texts.append(sentence)

        # Обрабатываем абзацы
        for paragraph in doc.paragraphs:
            # Пропускаем абзацы с изображениями
            has_image = any(
                'pic:pic' in run._element.xml
                for run in paragraph.runs
            )
            if has_image or not paragraph.text.strip():
                continue

            cleaned_text = self.preprocessor.remove_special_chars_docx(
                paragraph.text
            )
            if not cleaned_text.strip():
                continue

            new_para = new_doc.add_paragraph()
            self._copy_paragraph_format(paragraph, new_para)

            new_run = new_para.add_run(cleaned_text)
            if paragraph.runs:
                self._copy_run_format(paragraph.runs[0], new_run)

        # Добавляем текст из таблиц в конец
        for table_text in table_texts:
            new_doc.add_paragraph(table_text)

        new_doc.save(str(output_file))

    @staticmethod
    def _copy_paragraph_format(src, dst) -> None:
        """Копирует форматирование абзаца."""
        dst.style     = src.style
        dst.alignment = src.alignment
        fmt_src = src.paragraph_format
        fmt_dst = dst.paragraph_format
        fmt_dst.left_indent       = fmt_src.left_indent
        fmt_dst.right_indent      = fmt_src.right_indent
        fmt_dst.first_line_indent = fmt_src.first_line_indent
        fmt_dst.space_before      = fmt_src.space_before
        fmt_dst.space_after       = fmt_src.space_after
        fmt_dst.line_spacing      = fmt_src.line_spacing

    @staticmethod
    def _copy_run_format(src_run, dst_run) -> None:
        """Копирует форматирование run (жирный, курсив, шрифт и т.д.)."""
        dst_run.bold      = src_run.bold
        dst_run.italic    = src_run.italic
        dst_run.underline = src_run.underline
        if src_run.font.size:
            dst_run.font.size = src_run.font.size
        if src_run.font.name:
            dst_run.font.name = src_run.font.name
        try:
            if src_run.font.color.rgb:
                dst_run.font.color.rgb = src_run.font.color.rgb
        except Exception:
            pass